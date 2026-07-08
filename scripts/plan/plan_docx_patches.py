#!/usr/bin/env python3
"""生成 docx 补丁计划（表格定位，与 plan_sheet_patches 语义类似）。"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    raise SystemExit("请安装 python-docx: pip install python-docx") from e


def load_json(path: Path) -> dict:
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def week_matches(header: str, week: str, aliases: list[str]) -> bool:
    h = header.strip()
    for c in [week] + aliases:
        c = c.strip()
        if c == h or c in h or h in c:
            return True
    return False


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", type=Path, default=Path("config.json"))
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--extracted", type=Path, required=True)
    parser.add_argument("--output", type=Path, default=Path(".cache/patch-plan-docx.json"))
    args = parser.parse_args()

    cfg = load_json(args.config)
    extracted = load_json(args.extracted)
    by_name = {m["name"]: m["content"] for m in extracted.get("members", [])}

    doc = Document(args.input)
    plan = []

    for member in cfg.get("members", []):
        target = member.get("target", {})
        if target.get("type") != "docx_table_cell":
            plan.append({"name": member["name"], "status": "skip_type"})
            continue
        ti = int(target.get("table_index", 0))
        if ti >= len(doc.tables):
            plan.append({"name": member["name"], "status": "no_table"})
            continue
        table = doc.tables[ti]
        header_row = int(target.get("col_match", {}).get("header_row", 0))
        week = cfg.get("week", "")
        aliases = cfg.get("week_aliases", [])
        col_idx = None
        for j, cell in enumerate(table.rows[header_row].cells):
            if week_matches(cell.text, week, aliases):
                col_idx = j
                break
        name_cell = int(target.get("row_match", {}).get("cell", 0))
        contains = target.get("row_match", {}).get("contains", member["name"])
        row_idx = None
        for i, row in enumerate(table.rows):
            if i == header_row:
                continue
            if contains in row.cells[name_cell].text:
                row_idx = i
                break
        if row_idx is None or col_idx is None:
            plan.append({"name": member["name"], "status": "not_found"})
        else:
            plan.append(
                {
                    "name": member["name"],
                    "status": "ready",
                    "table_index": ti,
                    "row": row_idx,
                    "col": col_idx,
                    "content": by_name.get(member["name"], ""),
                }
            )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps({"patches": plan}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"patches": plan}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
